# -*- coding: utf-8 -*-
"""
DOCX Export Service – VCFE Database
Generates a professional Word document for a given CCCD profile.
"""

import io
import json
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from backend.constants import LOAI_HINH_DAC_THU


# ── Helpers ──────────────────────────────────────────────

def _safe(val) -> str:
    if val is None or str(val).strip() in ("", "N/A", "None", "nan"):
        return ""
    return str(val).strip()


def _fmt_date(val) -> str:
    s = _safe(val)
    if not s:
        return ""
    # Try to convert YYYY-MM-DD → DD/MM/YYYY
    for fmt_in, fmt_out in [("%Y-%m-%d", "%d/%m/%Y"), ("%d/%m/%Y", "%d/%m/%Y")]:
        try:
            return datetime.strptime(s, fmt_in).strftime(fmt_out)
        except ValueError:
            continue
    return s


def _set_tnr(run, size=12, bold=False):
    """Apply Times New Roman to a run."""
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def _heading_style(heading, size=14):
    for run in heading.runs:
        _set_tnr(run, size=size, bold=True)


def _add_kv(cell, key: str, value):
    val = _safe(value)
    if not val:
        return
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    k = p.add_run(f"{key}: ")
    k.bold = True
    p.add_run(val)


def _bullet(doc, bold_text: str, detail_text: str = "", note: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(bold_text)
    r1.bold = True
    if detail_text:
        p.add_run(detail_text)
    if note:
        rn = p.add_run(f"\n   Ghi chú: {note}")
        rn.italic = True


# ── Main Export ──────────────────────────────────────────

def generate_profile_docx(profile: dict, base_dir: str = "") -> Optional[bytes]:
    """
    Generate a DOCX report from a profile dict (from get_profile_full).
    Returns bytes of the .docx file, or None if profile is empty.
    """
    if not profile:
        return None

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Header ──
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_hdr.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    _set_tnr(r1, 13, bold=True)
    r2 = p_hdr.add_run("Độc lập - Tự do - Hạnh phúc")
    _set_tnr(r2, 12, bold=True)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = p_title.add_run("\nHỒ SƠ ĐỐI TƯỢNG CHI TIẾT\n")
    _set_tnr(rt, 16, bold=True)

    # ── 1. Thông tin cá nhân ──
    h1 = doc.add_heading("1. THÔNG TIN CÁ NHÂN", level=2)
    _heading_style(h1)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.columns[0].width = Inches(4.5)
    tbl.columns[1].width = Inches(1.5)
    cell_info = tbl.rows[0].cells[0]
    cell_img = tbl.rows[0].cells[1]

    _add_kv(cell_info, "Họ và tên", profile.get("ho_ten"))
    _add_kv(cell_info, "Số CCCD", profile.get("cccd"))
    _add_kv(cell_info, "Ngày sinh", _fmt_date(profile.get("ngay_sinh")))
    _add_kv(cell_info, "Giới tính", profile.get("gioi_tinh"))

    parts = [_safe(profile.get("dia_chi_xa")), _safe(profile.get("dia_chi_tinh"))]
    addr = " - ".join(p for p in parts if p)
    _add_kv(cell_info, "Thường trú", addr)
    _add_kv(cell_info, "Phân loại nghề nghiệp", profile.get("phan_loai_nghe_nghiep"))
    _add_kv(cell_info, "Chi tiết nơi làm việc", profile.get("chi_tiet_nghe_nghiep"))
    _add_kv(cell_info, "Ghi chú chung", profile.get("ghi_chu_chung"))

    # Avatar image
    avatar = profile.get("anh_chan_dung")
    if avatar and base_dir:
        avatar_path = Path(base_dir) / avatar
        if avatar_path.exists():
            p_img = cell_img.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                p_img.add_run().add_picture(str(avatar_path), width=Inches(1.38))
            except Exception:
                pass

    doc.add_paragraph()  # spacing

    # ── 2. Nhân thân ──
    nhan_than = profile.get("nhan_than", [])
    if nhan_than:
        h = doc.add_heading("2. THÔNG TIN THÂN NHÂN", level=2)
        _heading_style(h)
        for nt in nhan_than:
            details = []
            ns = _fmt_date(nt.get("ngay_sinh"))
            if ns:
                details.append(f"Sinh: {ns}")
            if _safe(nt.get("gioi_tinh")):
                details.append(f"Giới tính: {nt['gioi_tinh']}")
            if _safe(nt.get("nghe_nghiep")):
                details.append(f"Nghề nghiệp: {nt['nghe_nghiep']}")
            addr_parts = [_safe(nt.get("dia_chi_xa")), _safe(nt.get("dia_chi_tinh"))]
            addr_r = " - ".join(p for p in addr_parts if p)
            if addr_r:
                details.append(f"Địa chỉ: {addr_r}")
            detail_str = ""
            if details:
                detail_str = "\n   " + " | ".join(details)
            _bullet(doc,
                     f"{nt.get('quan_he', '')}: {nt.get('ho_ten', '')}",
                     detail_str,
                     _safe(nt.get("ghi_chu")))

    # ── 3. Liên hệ ──
    lien_he = profile.get("lien_he", [])
    if lien_he:
        h = doc.add_heading("3. LIÊN HỆ", level=2)
        _heading_style(h)
        for lh in lien_he:
            note = f" (Ghi chú: {lh['ghi_chu']})" if _safe(lh.get("ghi_chu")) else ""
            _bullet(doc, f"{lh.get('loai', '')}: ", _safe(lh.get("gia_tri")) + note)

    # ── 4. Tài khoản ngân hàng ──
    tai_chinh = profile.get("tai_chinh", [])
    if tai_chinh:
        h = doc.add_heading("4. TÀI KHOẢN NGÂN HÀNG", level=2)
        _heading_style(h)
        for tc in tai_chinh:
            val = _safe(tc.get("so_tai_khoan"))
            if _safe(tc.get("chu")):
                val += f" (Chủ TK: {tc['chu']})"
            _bullet(doc, f"{tc.get('ngan_hang', '')}: ", val, _safe(tc.get("ghi_chu")))

    # ── 5. Phương tiện ──
    phuong_tien = profile.get("phuong_tien", [])
    if phuong_tien:
        h = doc.add_heading("5. PHƯƠNG TIỆN", level=2)
        _heading_style(h)
        for pt in phuong_tien:
            val = _safe(pt.get("bien"))
            if _safe(pt.get("ten")):
                val += f" ({pt['ten']})"
            _bullet(doc, f"{pt.get('loai_xe', '')}: ", val, _safe(pt.get("ghi_chu")))

    # ── 6. Quá trình hoạt động ──
    qua_trinh = profile.get("qua_trinh", [])
    if qua_trinh:
        h = doc.add_heading("6. QUÁ TRÌNH HOẠT ĐỘNG", level=2)
        _heading_style(h)
        for qt in qua_trinh:
            _bullet(doc,
                     f"{_safe(qt.get('thoi_gian'))}\n",
                     f"   {_safe(qt.get('noi_dung'))}",
                     _safe(qt.get("ghi_chu")))

    # ── 7. Hồ sơ đặc thù CSXH ──
    ho_so = profile.get("ho_so_dac_thu", [])
    if ho_so:
        h = doc.add_heading("7. HỒ SƠ YẾU TỐ CSXH", level=2)
        _heading_style(h)
        for hs in ho_so:
            loai_key = hs.get("loai_hinh", "")
            loai_text = LOAI_HINH_DAC_THU.get(loai_key, loai_key)
            detail_lines = ""
            noi_dung_raw = hs.get("noi_dung", "")
            if noi_dung_raw:
                try:
                    parsed = json.loads(noi_dung_raw)
                    for k, v in parsed.items():
                        vs = _safe(v)
                        if vs:
                            label = k.replace("_", " ").title()
                            detail_lines += f"\n   {label}: {vs}"
                except (json.JSONDecodeError, TypeError):
                    if _safe(noi_dung_raw):
                        detail_lines = f"\n   Nội dung: {noi_dung_raw}"
            _bullet(doc, loai_text, detail_lines, _safe(hs.get("ghi_chu")))

    # ── Footer ──
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rf = p_footer.add_run(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    rf.italic = True
    rf.font.size = Pt(10)

    p_sys = doc.add_paragraph()
    p_sys.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rs = p_sys.add_run("VCFE Database")
    rs.font.size = Pt(9)
    rs.font.color.rgb = None  # default
    rs.italic = True

    # Save
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
